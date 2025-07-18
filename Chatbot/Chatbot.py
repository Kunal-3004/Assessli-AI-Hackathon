from threading import Thread
import traceback
from datetime import datetime
from enum import Enum
from typing import Dict, List, Optional, Any, Tuple, Union
from dataclasses import dataclass, field
import json
import time
from contextlib import contextmanager

from langchain_groq import ChatGroq
from langchain_openai import OpenAIEmbeddings
from langchain.document_loaders import WebBaseLoader, PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.pydantic_v1 import BaseModel, Field
from langchain import hub
from langchain_core.output_parsers import StrOutputParser
from langchain.prompts import PromptTemplate
import requests
from typing_extensions import TypedDict
from langgraph.graph import END, StateGraph, START
from duckduckgo_search import DDGS
from langchain_core.retrievers import BaseRetriever
from langchain.embeddings import HuggingFaceEmbeddings
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain.memory import ChatMessageHistory
from langchain.memory.chat_memory import BaseChatMessageHistory
from collections import defaultdict
from transformers import pipeline
import os
import tempfile
from werkzeug.utils import secure_filename

class SentimentLabel(Enum):
    POSITIVE = "positive"
    NEGATIVE = "negative"
    NEUTRAL = "neutral"
    MIXED = "mixed"
    
class ErrorType(Enum):
    RETRIEVAL_ERROR = "retrieval_error"
    GENERATION_ERROR = "generation_error"
    GRADING_ERROR = "grading_error"
    SEARCH_ERROR = "search_error"
    VALIDATION_ERROR = "validation_error"
    NETWORK_ERROR = "network_error"
    SYSTEM_ERROR = "system_error"


class ProcessingState(Enum):
    INITIALIZED = "initialized"
    RETRIEVING = "retrieving"
    GRADING = "grading"
    GENERATING = "generating"
    TRANSFORMING = "transforming"
    SEARCHING = "searching"
    COMPLETED = "completed"
    FAILED = "failed"
    RETRYING = "retrying"

@dataclass
class SentimentAnalysisResult:
    label: SentimentLabel
    score: float
    timestamp: datetime
    text_sample: str
    analysis_model: str

@dataclass
class SystemError:
    error_type: ErrorType
    message: str
    timestamp: datetime
    node: str
    retry_count: int = 0
    recoverable: bool = True


@dataclass
class ResponseResult:
    success: bool
    response: str
    error_message: Optional[str] = None
    state: ProcessingState = ProcessingState.COMPLETED
    retry_count: int = 0
    context_used: bool = False
    fallback_used: bool = False


class StateManager:
    
    def __init__(self):
        self.session_states: Dict[str, Dict] = {}
        self.max_retry_attempts = 3
        
    def initialize_session(self, session_id: str) -> Dict:
        if session_id not in self.session_states:
            self.session_states[session_id] = {
                'state': ProcessingState.INITIALIZED,
                'errors': [],
                'retry_count': 0,
                'start_time': datetime.now(),
                'last_activity': datetime.now(),
                'context': {},
                'consecutive_failures': 0
            }
        return self.session_states[session_id]
    
    def update_state(self, session_id: str, state: ProcessingState, context: Dict = None):
        session_state = self.initialize_session(session_id)
        session_state['state'] = state
        session_state['last_activity'] = datetime.now()
        if context:
            session_state['context'].update(context)
    
    def log_error(self, session_id: str, error: SystemError):
        session_state = self.initialize_session(session_id)
        session_state['errors'].append(error)
        session_state['consecutive_failures'] += 1
        
        # Keep only recent errors (last 5)
        if len(session_state['errors']) > 5:
            session_state['errors'] = session_state['errors'][-5:]
    
    def should_retry(self, session_id: str, error_type: ErrorType) -> bool:
        session_state = self.session_states.get(session_id, {})
        retry_count = session_state.get('retry_count', 0)
        consecutive_failures = session_state.get('consecutive_failures', 0)
        
        if error_type == ErrorType.VALIDATION_ERROR or consecutive_failures >= 3:
            return False
        
        return retry_count < self.max_retry_attempts
    
    def increment_retry(self, session_id: str):
        session_state = self.initialize_session(session_id)
        session_state['retry_count'] += 1
    
    def reset_failures(self, session_id: str):
        session_state = self.initialize_session(session_id)
        session_state['consecutive_failures'] = 0
        session_state['retry_count'] = 0
    
    def get_session_health(self, session_id: str) -> Dict:
        """Check session health - MISSING METHOD ADDED"""
        session_state = self.session_states.get(session_id, {})
        consecutive_failures = session_state.get('consecutive_failures', 0)
        
        return {
            'healthy': consecutive_failures < 5,
            'consecutive_failures': consecutive_failures,
            'retry_count': session_state.get('retry_count', 0)
        }
        
    def get_sentiment_summary(self, session_id: str) -> Dict:
        session_state = self.session_states.get(session_id, {})
        sentiment_history = session_state.get('context', {}).get('sentiment_history', [])
        
        if not sentiment_history:
            return {
                'total_analyses': 0,
                'positive_count': 0,
                'negative_count': 0,
                'neutral_count': 0,
                'average_score': 0.5,
                'dominant_sentiment': 'neutral'
            }
        
        positive_count = sum(1 for r in sentiment_history if r.label == SentimentLabel.POSITIVE)
        negative_count = sum(1 for r in sentiment_history if r.label == SentimentLabel.NEGATIVE)
        neutral_count = sum(1 for r in sentiment_history if r.label == SentimentLabel.NEUTRAL)
        average_score = sum(r.score for r in sentiment_history) / len(sentiment_history)
        
        dominant_sentiment = max(
            ('positive', positive_count),
            ('negative', negative_count),
            ('neutral', neutral_count),
            key=lambda x: x[1]
        )[0]
        
        return {
            'total_analyses': len(sentiment_history),
            'positive_count': positive_count,
            'negative_count': negative_count,
            'neutral_count': neutral_count,
            'average_score': average_score,
            'dominant_sentiment': dominant_sentiment
        }

class MultiModalProcessor:
    def __init__(self):
        self.supported_formats = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'csv': self._process_csv,
            'json': self._process_json,
            'xml': self._process_xml,
            'image': self._process_image,
            'video': self._process_video,
            'audio': self._process_audio
        }
        self.content_cache = {}
        
    def process_content(self, content_path: str, content_type: str) -> Dict:
        """Process multi-modal content and extract structured information"""
        cache_key = f"{content_path}_{content_type}"
        
        if cache_key in self.content_cache:
            return self.content_cache[cache_key]
            
        if content_type in self.supported_formats:
            result = self.supported_formats[content_type](content_path)
            self.content_cache[cache_key] = result
            return result
        else:
            return {"error": f"Unsupported content type: {content_type}"}
            
    def _process_pdf(self, path: str) -> Dict:
        """Process PDF and extract text, tables, metadata"""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(path)
            
            content = {
                'text': '',
                'metadata': reader.metadata,
                'pages': len(reader.pages),
                'tables': [],
                'images': []
            }
            
            for page in reader.pages:
                content['text'] += page.extract_text()
                
            return content
        except Exception as e:
            return {"error": f"PDF processing failed: {str(e)}"}
            
    def _process_docx(self, path: str) -> Dict:
        """Process Word document"""
        try:
            from docx import Document
            doc = Document(path)
            
            content = {
                'text': '',
                'paragraphs': [],
                'tables': [],
                'images': []
            }
            
            for paragraph in doc.paragraphs:
                content['text'] += paragraph.text + '\n'
                content['paragraphs'].append(paragraph.text)
                
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                content['tables'].append(table_data)
                
            return content
        except Exception as e:
            return {"error": f"DOCX processing failed: {str(e)}"}
            
    def _process_csv(self, path: str) -> Dict:
        """Process CSV file"""
        try:
            import pandas as pd
            df = pd.read_csv(path)
            
            content = {
                'data': df.to_dict('records'),
                'columns': df.columns.tolist(),
                'shape': df.shape,
                'summary': df.describe().to_dict(),
                'sample': df.head().to_dict('records')
            }
            
            return content
        except Exception as e:
            return {"error": f"CSV processing failed: {str(e)}"}
            
    def _process_json(self, path: str) -> Dict:
        """Process JSON file"""
        try:
            with open(path, 'r') as f:
                data = json.load(f)
                
            content = {
                'data': data,
                'structure': self._analyze_json_structure(data),
                'size': len(str(data))
            }
            
            return content
        except Exception as e:
            return {"error": f"JSON processing failed: {str(e)}"}
            
    def _process_xml(self, path: str) -> Dict:
        """Process XML file"""
        try:
            import xml.etree.ElementTree as ET
            tree = ET.parse(path)
            root = tree.getroot()
            
            content = {
                'root_tag': root.tag,
                'structure': self._xml_to_dict(root),
                'namespaces': dict(root.nsmap) if hasattr(root, 'nsmap') else {}
            }
            
            return content
        except Exception as e:
            return {"error": f"XML processing failed: {str(e)}"}
            
    def _process_image(self, path: str) -> Dict:
        """Process image and extract text/metadata"""
        try:
            from PIL import Image
            import pytesseract
            
            img = Image.open(path)
            
            content = {
                'text': pytesseract.image_to_string(img),
                'size': img.size,
                'format': img.format,
                'mode': img.mode,
                'has_text': bool(pytesseract.image_to_string(img).strip())
            }
            
            return content
        except Exception as e:
            return {"error": f"Image processing failed: {str(e)}"}
            
    def _process_video(self, path: str) -> Dict:
        """Process video and extract metadata/transcript"""
        try:
            import cv2
            
            cap = cv2.VideoCapture(path)
            
            content = {
                'duration': cap.get(cv2.CAP_PROP_FRAME_COUNT) / cap.get(cv2.CAP_PROP_FPS),
                'fps': cap.get(cv2.CAP_PROP_FPS),
                'width': int(cap.get(cv2.CAP_PROP_FRAME_WIDTH)),
                'height': int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT)),
                'frame_count': int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            }
            
            cap.release()
            return content
        except Exception as e:
            return {"error": f"Video processing failed: {str(e)}"}
            
    def _process_audio(self, path: str) -> Dict:
        """Process audio and extract transcript"""
        try:
            import librosa
            
            audio, sr = librosa.load(path)
            
            content = {
                'duration': len(audio) / sr,
                'sample_rate': sr,
                'channels': 1 if audio.ndim == 1 else audio.shape[0],
                'audio_features': {
                    'tempo': librosa.beat.tempo(audio, sr=sr)[0],
                    'spectral_centroid': float(librosa.feature.spectral_centroid(audio, sr=sr).mean())
                }
            }
            
            return content
        except Exception as e:
            return {"error": f"Audio processing failed: {str(e)}"}
            
    def _analyze_json_structure(self, data: Any, depth: int = 0) -> Dict:
        """Analyze JSON structure"""
        if depth > 3:  
            return {"type": "truncated"}
            
        if isinstance(data, dict):
            return {
                "type": "object",
                "keys": list(data.keys()),
                "structure": {k: self._analyze_json_structure(v, depth + 1) for k, v in data.items()}
            }
        elif isinstance(data, list):
            return {
                "type": "array",
                "length": len(data),
                "item_type": self._analyze_json_structure(data[0], depth + 1) if data else "empty"
            }
        else:
            return {"type": type(data).__name__}
            
    def _xml_to_dict(self, element) -> Dict:
        """Convert XML element to dictionary"""
        result = {}
        if element.text and element.text.strip():
            result['text'] = element.text.strip()
        
        for child in element:
            child_data = self._xml_to_dict(child)
            if child.tag in result:
                if not isinstance(result[child.tag], list):
                    result[child.tag] = [result[child.tag]]
                result[child.tag].append(child_data)
            else:
                result[child.tag] = child_data
                
        return result

class AgenticRAG:
    
    def __init__(self):
        self.state_manager = StateManager()
        self.setup_system()
        
        self.multimodal_processor = MultiModalProcessor()
    
    def setup_system(self):
        try:
            # Set GROQ API key
            os.environ["GROQ_API_KEY"] = "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"
            
            # Initialize LLM
            self.llm = ChatGroq(model_name="llama-3.1-8b-instant", temperature=0)
            
            # Initialize memory
            self.chat_histories = defaultdict(ChatMessageHistory)
            
            # Setup URLs and documents
            self.setup_documents()
            
            # Setup retrieval components
            self.setup_retrieval()
            
            # Setup grading components
            self.setup_grading()
            
            # Setup workflow
            self.setup_workflow()
            
            self.setup_sentiment_analysis()
            
            print("System setup completed successfully!")
            
        except Exception as e:
            print(f"System initialization failed: {str(e)}")
            print(f"Error traceback: {traceback.format_exc()}")
            raise
    
    def setup_documents(self):
        urls = [
            "https://www.assessli.com/", 
            "https://www.assessli.com/contactus",
            "https://www.youtube.com/@assessli",
            "https://www.linkedin.com/company/assessli"
        ]
        
        try:
            # Load documents
            docs = []
            for url in urls:
                try:
                    print(f"Loading URL: {url}")
                    loader = WebBaseLoader(url)
                    loaded_docs = loader.load()
                    for doc in loaded_docs:
                        doc.metadata["source"] = url
                    docs.extend(loaded_docs)
                    print(f"Loaded {len(loaded_docs)} documents from {url}")
                except Exception as e:
                    print(f"Failed to load {url}: {str(e)}")
                    continue
            
            if not docs:
                print("No documents could be loaded from URLs")
                # Create a fallback document with basic info
                fallback_doc = Document(
                    page_content="Assessli is a company that provides assessment solutions. For more information, visit their website at assessli.com or contact them through their contact page.",
                    metadata={"source": "fallback"}
                )
                docs = [fallback_doc]
                print("Created fallback document")
            
            # Split documents
            text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
                chunk_size=250, chunk_overlap=0
            )
            self.doc_splits = text_splitter.split_documents(docs)
            print(f"Split documents into {len(self.doc_splits)} chunks")
            
            # Create vector store
            embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
            self.vectorstore = FAISS.from_documents(
                documents=self.doc_splits,
                embedding=embedding_model,
            )
            print("Created vector store")
            
        except Exception as e:
            print(f"Document setup failed: {str(e)}")
            print(f"Error traceback: {traceback.format_exc()}")
            raise Exception(f"Document setup failed: {str(e)}")
        
    def setup_sentiment_analysis(self):
        try:
            self.sentiment_analyzer = pipeline(
                "sentiment-analysis",
                model="distilbert-base-uncased-finetuned-sst-2-english",
                framework="pt",
                device=-1  
            )
            print("Sentiment analysis pipeline initialized")
        except Exception as e:
            print(f"Failed to initialize sentiment analyzer: {str(e)}")
            self.sentiment_analyzer = self.analyzer
            
    def analyzer(self, text: str) -> dict:
        """Fallback simple sentiment analyzer"""
        positive_words = {"good", "great", "excellent", "happy", "positive"}
        negative_words = {"bad", "poor", "terrible", "unhappy", "negative", "slow", "disappoint"}
        
        words = set(text.lower().split())
        pos_count = len(words & positive_words)
        neg_count = len(words & negative_words)
        
        if pos_count > neg_count:
            return {"label": "POSITIVE", "score": min(0.9, pos_count * 0.3)}
        elif neg_count > pos_count:
            return {"label": "NEGATIVE", "score": min(0.9, neg_count * 0.3)}
        else:
            return {"label": "NEUTRAL", "score": 0.5}
        
    def analyze_sentiment(self, text: str, session_id: str = "default") -> SentimentAnalysisResult:
        try:
            start_time = time.time()
            
            analysis_text = text[:2000] if len(text) > 2000 else text
            
            result = self.sentiment_analyzer(analysis_text)
            
            if isinstance(result, list):
                result = result[0]  
            
            label = SentimentLabel(result['label'].lower())
            score = float(result['score'])
            
            analysis_result = SentimentAnalysisResult(
                label=label,
                score=score,
                timestamp=datetime.now(),
                text_sample=analysis_text[:100] + "..." if len(analysis_text) > 100 else analysis_text,
                analysis_model=str(self.sentiment_analyzer.model.__class__.__name__)
            )
            
            session_state = self.state_manager.initialize_session(session_id)
            if 'sentiment_history' not in session_state['context']:
                session_state['context']['sentiment_history'] = []
            session_state['context']['sentiment_history'].append(analysis_result)
            
            if len(session_state['context']['sentiment_history']) > 10:
                session_state['context']['sentiment_history'] = session_state['context']['sentiment_history'][-10:]
            
            print(f"Sentiment analysis completed in {time.time() - start_time:.2f}s - {label.value} ({score:.2f})")
            return analysis_result
            
        except Exception as e:
            print(f"Sentiment analysis failed: {str(e)}")
            error = SystemError(
                error_type=ErrorType.SYSTEM_ERROR,
                message=str(e),
                timestamp=datetime.now(),
                node="sentiment_analysis"
            )
            self.state_manager.log_error(session_id, error)
            
            return SentimentAnalysisResult(
                label=SentimentLabel.NEUTRAL,
                score=0.5,
                timestamp=datetime.now(),
                text_sample="",
                analysis_model="fallback"
            )

    def get_sentiment_trend(self, session_id: str) -> Tuple[SentimentLabel, float]:
        try:
            session_state = self.state_manager.initialize_session(session_id)
            history = session_state['context'].get('sentiment_history', [])
            
            if not history:
                return SentimentLabel.NEUTRAL, 0.5
                
            avg_score = sum(r.score for r in history) / len(history)
            
            pos_count = sum(1 for r in history if r.label == SentimentLabel.POSITIVE)
            neg_count = sum(1 for r in history if r.label == SentimentLabel.NEGATIVE)
            
            if pos_count > neg_count:
                overall_label = SentimentLabel.POSITIVE
            elif neg_count > pos_count:
                overall_label = SentimentLabel.NEGATIVE
            else:
                overall_label = SentimentLabel.NEUTRAL
                
            return overall_label, avg_score
            
        except Exception as e:
            print(f"Failed to calculate sentiment trend: {str(e)}")
            return SentimentLabel.NEUTRAL, 0.5

    
    def setup_retrieval(self):
        """Setup retrieval components"""
        try:
            # Define relevance score model
            class RelevantScore(BaseModel):
                score: float = Field(description="The relevance score of the document to the query", example=8.0)
            
            self.RelevantScore = RelevantScore
            
            # Setup retriever wrapper
            class PydanticAdaptiveRetriever(BaseRetriever):
                adaptive_retriever: callable = Field(exclude=True)
                
                class Config:
                    arbitrary_types_allowed = True
                
                def get_relevant_documents(self, query: str) -> List[Document]:
                    return self.adaptive_retriever(query)
                
                async def aget_relevant_documents(self, query: str) -> List[Document]:
                    return self.get_relevant_documents(query)
            
            self.retriever = PydanticAdaptiveRetriever(adaptive_retriever=self.adaptive_retrieve)
            print("Retrieval components setup completed")
            
        except Exception as e:
            print(f"Retrieval setup failed: {str(e)}")
            raise
    
    def setup_grading(self):
        """Setup document and response grading"""
        try:
            # Document grading
            class GradeDocuments(BaseModel):
                binary_score: str = Field(description="Documents are relevant to the question, 'yes' or 'no'")
            
            # Hallucination grading
            class GradeHallucinations(BaseModel):
                binary_score: str = Field(description="Answer is grounded in the facts, 'yes' or 'no'")
            
            # Answer grading
            class GradeAnswer(BaseModel):
                binary_score: str = Field(description="Answer addresses the question, 'yes' or 'no'")
            
            # Setup grading chains
            self.setup_grading_chains(GradeDocuments, GradeHallucinations, GradeAnswer)
            print("Grading components setup completed")
            
        except Exception as e:
            print(f"Grading setup failed: {str(e)}")
            raise
    
    def setup_grading_chains(self, GradeDocuments, GradeHallucinations, GradeAnswer):
        """Setup grading chain components"""
        try:
            # Document relevance grader
            system = """You are a grader assessing relevance of a retrieved document to a user question. 
            If the document contains keyword(s) or semantic meaning related to the user question, grade it as relevant. 
            Give a binary score 'yes' or 'no' score to indicate whether the document is relevant to the question."""
            
            grade_prompt = ChatPromptTemplate.from_messages([
                ("system", system),
                ("human", "Retrieved document: \n\n {document} \n\n User question: {question}")
            ])
            
            self.retrieval_grader = grade_prompt | self.llm.with_structured_output(GradeDocuments)
            
            # Hallucination grader
            system = """You are a grader assessing whether an LLM generation is grounded in / supported by a set of retrieved facts. 
            Give a binary score 'yes' or 'no'. 'Yes' means that the answer is grounded in / supported by the set of facts."""
            
            hallucination_prompt = ChatPromptTemplate.from_messages([
                ("system", system),
                ("human", "Set of facts: \n\n {documents} \n\n LLM generation: {generation}"),
            ])
            
            self.hallucination_grader = hallucination_prompt | self.llm.with_structured_output(GradeHallucinations)
            
            # Answer grader
            system = """You are a grader assessing whether an answer addresses / resolves a question 
            Give a binary score 'yes' or 'no'. Yes' means that the answer resolves the question."""
            
            answer_prompt = ChatPromptTemplate.from_messages([
                ("system", system),
                ("human", "User question: \n\n {question} \n\n LLM generation: {generation}"),
            ])
            
            self.answer_grader = answer_prompt | self.llm.with_structured_output(GradeAnswer)
            
            # Question rewriter
            system = """You a question re-writer that converts an input question to a better version that is optimized 
            for vectorstore retrieval. Look at the input and try to reason about the underlying semantic intent / meaning."""
            
            re_write_prompt = ChatPromptTemplate.from_messages([
                ("system", system),
                ("human", "Here is the initial question: \n\n {question} \n Formulate an improved question."),
            ])
            
            self.question_rewriter = re_write_prompt | self.llm | StrOutputParser()
            
        except Exception as e:
            print(f"Grading chains setup failed: {str(e)}")
            raise Exception(f"Grading setup failed: {str(e)}")
    
    def setup_workflow(self):
        """Setup the workflow graph"""
        try:
            # Graph state
            class GraphState(TypedDict):
                question: str
                generation: str
                documents: List[str]
                session_id: str
                error_message: str
                retry_count: int
            
            self.GraphState = GraphState
            
            # Create workflow
            workflow = StateGraph(GraphState)
            
            # Add nodes
            workflow.add_node("retrieve", self.safe_retrieve)
            workflow.add_node("grade_documents", self.safe_grade_documents)
            workflow.add_node("generate", self.safe_generate)
            workflow.add_node("transform_query", self.safe_transform_query)
            
            # Add edges
            workflow.add_edge(START, "retrieve")
            workflow.add_edge("retrieve", "grade_documents")
            workflow.add_conditional_edges(
                "grade_documents",
                self.decide_to_generate,
                {
                    "transform_query": "transform_query",
                    "generate": "generate",
                },
            )
            workflow.add_edge("transform_query", "retrieve")
            workflow.add_conditional_edges(
                "generate",
                self.grade_generation_v_documents_and_question,
                {
                    "not supported": "generate",
                    "useful": END,
                    "not useful": "transform_query",
                },
            )
            
            self.app = workflow.compile()
            print("Workflow setup completed")
            
        except Exception as e:
            print(f"Workflow setup failed: {str(e)}")
            raise
    
    def safe_operation(self, operation_name: str, operation_func, *args, **kwargs):
        """Wrapper for safe operation execution"""
        try:
            return operation_func(*args, **kwargs)
        except Exception as e:
            error_message = f"Error in {operation_name}: {str(e)}"
            print(f"{error_message}")
            return {"error": error_message, "success": False}
    
    def classify_query(self, query: str) -> str:
        """Classify query type with error handling"""
        try:
            prompt = PromptTemplate(
                input_variables=["query"],
                template=(
                    "Classify the following query into one of these categories: "
                    "Factual, Analytical, Opinion, or Contextual.\n\n"
                    "Query: {query}\n\n"
                    "Category:"
                )
            )
            chain = prompt | self.llm | StrOutputParser()
            result = chain.invoke({"query": query})
            return result.strip()
        except Exception as e:
            print(f"Query classification failed: {str(e)}")
            return "Factual"  # Default classification
        
    class SelectedIndices(BaseModel):
            indices: List[int] = Field(description="Indices of selected documents", example=[0, 1, 2, 3])

    class SubQueries(BaseModel):
        sub_queries: List[str] = Field(description="List of sub-queries for comprehensive analysis", example=["What is the population of New York?", "What is the GDP of New York?"])

    def retrieve_factual(self, query: str, k: int = 4) -> List[Document]:
        """Retrieve factual documents with error handling"""
        try:
            # Enhance query
            enhance_prompt = PromptTemplate(
                input_variables=["query"],
                template="Enhance this factual query for better information retrieval: {query}"
            )
            query_chain = enhance_prompt | self.llm
            enhanced_query = query_chain.invoke({"query": query}).content
            
            # Retrieve documents
            docs = self.vectorstore.similarity_search(enhanced_query, k=k*2)
            
            # Simple ranking fallback if structured output fails
            try:
                # Rank documents
                ranking_prompt = PromptTemplate(
                    input_variables=["query", "doc"],
                    template="On a scale of 1-10, how relevant is this document to the query: '{query}'?\nDocument: {doc}\nRelevance score:"
                )
                ranking_chain = ranking_prompt | self.llm.with_structured_output(self.RelevantScore)
                
                ranked_docs = []
                for doc in docs:
                    try:
                        input_data = {"query": enhanced_query, "doc": doc.page_content}
                        score = float(ranking_chain.invoke(input_data).score)
                        ranked_docs.append((doc, score))
                    except Exception:
                        ranked_docs.append((doc, 5.0))  # Default score
                
                ranked_docs.sort(key=lambda x: x[1], reverse=True)
                return [doc for doc, _ in ranked_docs[:k]]
            except Exception as e:
                print(f"Document ranking failed: {str(e)}")
                return docs[:k]
            
        except Exception as e:
            print(f"Factual retrieval failed: {str(e)}")
            # Fallback to basic similarity search
            return self.vectorstore.similarity_search(query, k=k)
        
    def retrieve_analytical(self, query: str, k: int = 4) -> List[Document]:
        print("Retrieving Analytical Documents...")
        
        try:
            sub_queries_prompt = PromptTemplate(
                input_variables=["query", "k"],
                template="Generate {k} sub-questions for: {query}"
            )
            sub_queries_chain = sub_queries_prompt | self.llm.with_structured_output(self.SubQueries)
            sub_queries = sub_queries_chain.invoke({"query": query, "k": k}).sub_queries
            print(f"Generated Sub-queries: {sub_queries}")

            all_docs = []
            for sq in sub_queries:
                all_docs.extend(self.vectorstore.similarity_search(sq, k=2))

            docs_text = "\n".join([f"{i}: {doc.page_content[:50]}..." for i, doc in enumerate(all_docs)])

            diversity_prompt = PromptTemplate(
                input_variables=["query", "docs", "k"],
                template="Select the most diverse and relevant set of {k} documents for the query: '{query}'\nDocuments: {docs}\nReturn only the indices of selected documents as a list of integers."
            )
            diversity_chain = diversity_prompt | self.llm.with_structured_output(self.SelectedIndices)

            selected_indices = diversity_chain.invoke({"query": query, "docs": docs_text, "k": k}).indices
            return [all_docs[i] for i in selected_indices if i < len(all_docs)]

        except Exception as e:
            print(f"Analytical retrieval error: {e}")
            return self.vectorstore.similarity_search(query, k=k)

    def retrieve_opinion(self, query: str, k: int = 3) -> List[Document]:
        print("Retrieving Opinion-based Documents...")

        try:
            opinion_prompt = PromptTemplate(
                input_variables=["query", "k"],
                template="Identify {k} distinct viewpoints or perspectives on the topic: {query}"
            )
            viewpoints = (opinion_prompt | self.llm).invoke({"query": query, "k": k}).content.split('\n')
            print(f"Identified Viewpoints: {viewpoints}")

            all_docs = []
            for vp in viewpoints:
                all_docs.extend(self.vectorstore.similarity_search(f"{query} {vp}", k=2))

            docs_text = "\n".join([f"{i}: {doc.page_content[:100]}..." for i, doc in enumerate(all_docs)])

            opinion_select_prompt = PromptTemplate(
                input_variables=["query", "docs", "k"],
                template="Classify these documents into distinct opinions on '{query}' and select the {k} most representative and diverse viewpoints:\nDocuments: {docs}\nSelected indices:"
            )
            opinion_chain = opinion_select_prompt | self.llm.with_structured_output(self.SelectedIndices)

            selected_indices = opinion_chain.invoke({"query": query, "docs": docs_text, "k": k}).indices
            return [all_docs[i] for i in selected_indices if i < len(all_docs)]

        except Exception as e:
            print(f"Opinion retrieval error: {e}")
            return self.vectorstore.similarity_search(query, k=k)
        
    def retrieve_contextual(self, query: str, k: int = 4, user_context: str = None) -> List[Document]:
        print("Retrieving Contextual Documents...")

        try:
            context_str = user_context or "No specific context provided"

            contextualize_prompt = PromptTemplate(
                input_variables=["query", "context"],
                template="Given the user context: {context}\nReformulate the query to best address the user's needs: {query}"
            )
            contextualized_query = (contextualize_prompt | self.llm).invoke({"query": query, "context": context_str}).content
            print(f"Contextualized Query: {contextualized_query}")

            docs = self.vectorstore.similarity_search(contextualized_query, k=k*2)

            contextual_rank_prompt = PromptTemplate(
                input_variables=["query", "context", "doc"],
                template="Given the query: '{query}' and user context: '{context}', rate the relevance of this document on a scale of 1-10:\nDocument: {doc}\nRelevance score:"
            )
            contextual_rank_chain = contextual_rank_prompt | self.llm.with_structured_output(self.RelevantScore)

            ranked_docs = []
            for doc in docs:
                input_data = {"query": contextualized_query, "context": context_str, "doc": doc.page_content}
                score = float(contextual_rank_chain.invoke(input_data).score)
                ranked_docs.append((doc, score))

            ranked_docs.sort(key=lambda x: x[1], reverse=True)
            return [doc for doc, _ in ranked_docs[:k]]

        except Exception as e:
            print(f"Contextual retrieval error: {e}")
            return self.vectorstore.similarity_search(query, k=k)


    
    def adaptive_retrieve(self, query: str, k: int = 4, user_context: str = None) -> List[Document]:
        """Adaptive retrieval based on query classification"""
        try:
            category = self.classify_query(query)
            print(f"Query category classified as: {category}")

            if category == "Factual":
                return self.retrieve_factual(query, k)
            elif category == "Analytical":
                return self.retrieve_analytical(query, k)
            elif category == "Opinion":
                return self.retrieve_opinion(query, k)
            elif category == "Contextual":
                return self.retrieve_contextual(query, k, user_context)
            else:
                # Unknown or unsupported category fallback
                print(f"Unrecognized category '{category}', using basic similarity search.")
                return self.vectorstore.similarity_search(query, k)

        except Exception as e:
            print(f"Adaptive retrieval failed: {str(e)}")
            try:
                return self.vectorstore.similarity_search(query, k)
            except Exception as e2:
                print(f"Basic retrieval also failed: {str(e2)}")
                return []

    
    def duckduckgo_search(self, query, max_results=3):
        """DuckDuckGo search with error handling"""
        try:
            print(f"🔍 Searching web for: {query}")
            time.sleep(1)
            with DDGS() as ddgs:
                results = ddgs.text(query, max_results=max_results)
                docs = []
                for r in results:
                    if 'body' in r:
                        docs.append(Document(page_content=r['body'], metadata={"source": r['href']}))
                        if len(docs) >= max_results:
                            break
                print(f"Found {len(docs)} web results")
                return docs
        except Exception as e:
            print(f"Web search failed: {str(e)}")
            return []
    
    def safe_retrieve(self, state):
        """Safe document retrieval"""
        try:
            question = state["question"]
            session_id = state.get("session_id", "default")
            
            print(f"Retrieving documents for: {question}")
            self.state_manager.update_state(session_id, ProcessingState.RETRIEVING)
            
            documents = self.retriever.invoke(question)
            
            if not documents:
                print("No documents found, will trigger search fallback")
                return {"documents": [], "question": question, "session_id": session_id, 
                       "error_message": "No documents found for your query."}
            
            print(f"Retrieved {len(documents)} documents")
            return {"documents": documents, "question": question, "session_id": session_id}
            
        except Exception as e:
            print(f"Retrieval failed: {str(e)}")
            session_id = state.get("session_id", "default")
            error = SystemError(
                error_type=ErrorType.RETRIEVAL_ERROR,
                message=str(e),
                timestamp=datetime.now(),
                node="retrieve"
            )
            self.state_manager.log_error(session_id, error)
            
            return {"documents": [], "question": state["question"], "session_id": session_id,
                   "error_message": "I'm having trouble finding relevant information. Let me try a different approach."}
    
    def safe_grade_documents(self, state):
        """Safe document grading"""
        try:
            question = state["question"]
            documents = state["documents"]
            session_id = state.get("session_id", "default")
            
            print(f"Grading {len(documents)} documents")
            self.state_manager.update_state(session_id, ProcessingState.GRADING)
            
            if not documents:
                print("No documents found, trying web search")
                # Try web search as fallback
                search_docs = self.duckduckgo_search(question)
                if search_docs:
                    print(f"Web search found {len(search_docs)} documents")
                    return {"documents": search_docs, "question": question, "session_id": session_id}
                else:
                    print("Web search also failed")
                    return {"documents": [], "question": question, "session_id": session_id,
                           "error_message": "I couldn't find relevant information for your question."}
            
            filtered_docs = []
            for i, d in enumerate(documents):
                try:
                    score = self.retrieval_grader.invoke(
                        {"question": question, "document": d.page_content}
                    )
                    if score.binary_score == "yes":
                        filtered_docs.append(d)
                        print(f"Document {i+1} is relevant")
                    else:
                        print(f"Document {i+1} is not relevant")
                except Exception as e:
                    print(f"Failed to grade document {i+1}: {str(e)}")
                    # Include document if grading fails
                    filtered_docs.append(d)
            
            # If no relevant docs found, try web search
            if not filtered_docs:
                print("🔍 No relevant docs found, trying web search")
                search_docs = self.duckduckgo_search(question)
                filtered_docs = search_docs if search_docs else documents[:2]  # Use some docs as fallback
            
            print(f"Final document count: {len(filtered_docs)}")
            return {"documents": filtered_docs, "question": question, "session_id": session_id}
            
        except Exception as e:
            print(f"Document grading failed: {str(e)}")
            session_id = state.get("session_id", "default")
            error = SystemError(
                error_type=ErrorType.GRADING_ERROR,
                message=str(e),
                timestamp=datetime.now(),
                node="grade_documents"
            )
            self.state_manager.log_error(session_id, error)
            
            # Return original documents as fallback
            return {"documents": state.get("documents", []), "question": question, "session_id": session_id}
    
    def safe_generate(self, state):
        """Safe response generation"""
        try:
            question = state["question"]
            documents = state["documents"]
            session_id = state.get("session_id", "default")
            
            print(f"Generating response using {len(documents)} documents")
            self.state_manager.update_state(session_id, ProcessingState.GENERATING)
            
            if not documents:
                print("No documents available for generation")
                return {"documents": documents, "question": question, "session_id": session_id,
                       "generation": "I apologize, but I don't have enough information to answer your question properly. Could you please rephrase your question or provide more context?"}
            
            try:
                # Setup RAG chain
                prompt = hub.pull("rlm/rag-prompt")
                rag_chain = prompt | self.llm | StrOutputParser()
                
                # Generate response
                generation = rag_chain.invoke({"context": documents, "question": question})
                
                if not generation or len(generation.strip()) < 10:
                    generation = "I found some relevant information but couldn't generate a complete response. Could you please ask your question in a different way?"
                
                print(f"Generated response: {generation[:100]}...")
                return {"documents": documents, "question": question, "generation": generation, "session_id": session_id}
                
            except Exception as e:
                print(f"RAG chain failed: {str(e)}")
                # Simple fallback generation
                context_text = "\n".join([doc.page_content for doc in documents[:3]])
                simple_prompt = f"Based on this context: {context_text}\n\nAnswer this question: {question}"
                
                try:
                    generation = self.llm.invoke(simple_prompt).content
                    print(f"Fallback generation successful")
                    return {"documents": documents, "question": question, "generation": generation, "session_id": session_id}
                except Exception as e2:
                    print(f"Fallback generation also failed: {str(e2)}")
                    raise e2
            
        except Exception as e:
            print(f"Generation completely failed: {str(e)}")
            session_id = state.get("session_id", "default")
            error = SystemError(
                error_type=ErrorType.GENERATION_ERROR,
                message=str(e),
                timestamp=datetime.now(),
                node="generate"
            )
            self.state_manager.log_error(session_id, error)
            
            fallback_response = "I'm experiencing some technical difficulties generating a response. Please try asking your question again, or rephrase it for better results."
            return {"documents": state.get("documents", []), "question": question, 
                   "generation": fallback_response, "session_id": session_id}
    
    def safe_transform_query(self, state):
        """Safe query transformation"""
        try:
            question = state["question"]
            session_id = state.get("session_id", "default")
            
            print(f"Transforming query: {question}")
            self.state_manager.update_state(session_id, ProcessingState.TRANSFORMING)
            
            better_question = self.question_rewriter.invoke({"question": question})
            
            if not better_question or len(better_question.strip()) < 5:
                better_question = question  # Use original question
            
            print(f"Transformed query: {better_question}")
            return {"documents": state.get("documents", []), "question": better_question, "session_id": session_id}
            
        except Exception as e:
            print(f"Query transformation failed: {str(e)}")
            session_id = state.get("session_id", "default")
            error = SystemError(
                error_type=ErrorType.SYSTEM_ERROR,
                message=str(e),
                timestamp=datetime.now(),
                node="transform_query"
            )
            self.state_manager.log_error(session_id, error)
            
            # Return original question
            return {"documents": state.get("documents", []), "question": question, "session_id": session_id}
    
    def decide_to_generate(self, state):
        """Decide whether to generate or transform query"""
        documents = state.get("documents", [])
        
        if not documents:
            print("No documents found, will transform query")
            return "transform_query"
        else:
            print("Documents found, proceeding to generate")
            return "generate"
    
    def grade_generation_v_documents_and_question(self, state):
        """Grade generation quality"""
        try:
            question = state["question"]
            documents = state["documents"]
            generation = state["generation"]
            session_id = state.get("session_id", "default")
            
            print("Grading generation quality")
            
            # Check for error messages in generation
            if "technical difficulties" in generation.lower() or "apologize" in generation.lower():
                print("Generation contains error messages")
                return "not useful"
            
            # Try to grade hallucinations
            try:
                score = self.hallucination_grader.invoke(
                    {"documents": documents, "generation": generation}
                )
                if score.binary_score == "no":
                    print("Generation not supported by documents")
                    return "not supported"
                else:
                    print("Generation is supported by documents")
            except Exception as e:
                print(f"Hallucination grading failed: {str(e)}")
                pass  # Continue if grading fails
            
            # Try to grade answer quality
            try:
                score = self.answer_grader.invoke({"question": question, "generation": generation})
                if score.binary_score == "no":
                    print("Generation doesn't address the question")
                    return "not useful"
                else:
                    print("Generation addresses the question")
            except Exception as e:
                print(f"Answer grading failed: {str(e)}")
                pass  # Continue if grading fails
            
            # If we reach here, consider it useful
            print("Generation is useful")
            self.state_manager.reset_failures(session_id)
            return "useful"
            
        except Exception as e:
            print(f"Generation grading failed: {str(e)}")
            return "useful"  # Default to useful if grading fails
    
    def generate_response(self, question: str, session_id: str = "default") -> str:
        """Generate response with comprehensive error handling"""
        try:
            print(f"Starting response generation for: {question}")
            
            # Initialize session
            self.state_manager.initialize_session(session_id)
            
            # Validate input
            if not question or len(question.strip()) < 3:
                print("Question too short")
                return "Please provide a more specific question so I can help you better."
            
            # Check session health
            health = self.state_manager.get_session_health(session_id)
            if not health['healthy']:
                print("Session unhealthy")
                return "I'm experiencing some issues. Please try starting a new conversation."
            
            # Process the question
            inputs = {"question": question, "session_id": session_id}
            
            try:
                print("Running workflow...")
                # Run the workflow
                final_output = None
                step_count = 0
                
                for output in self.app.stream(inputs):
                    step_count += 1
                    print(f"Step {step_count}: {list(output.keys())}")
                    
                    for key, value in output.items():
                        if key == "generate":
                            final_output = value
                            print(f"Found final output in step {step_count}")
                            break
                    
                    if final_output:
                        break
                    
                    # Safety check to prevent infinite loops
                    if step_count > 10:
                        print("Too many workflow steps, breaking")
                        break
                
                if final_output and "generation" in final_output:
                    response = final_output["generation"]
                    
                    # Validate response quality
                    if len(response.strip()) < 10:
                        print("Response too short")
                        return "I found some information but couldn't provide a complete answer. Could you please rephrase your question?"
                    
                    print("Response generation successful")
                    self.state_manager.update_state(session_id, ProcessingState.COMPLETED)
                    return response
                else:
                    print("No generation found in workflow output")
                    return "I'm having trouble processing your question right now. Please try rephrasing it or ask something else."
                    
            except Exception as e:
                print(f"Workflow execution failed: {str(e)}")
                print(f"Workflow traceback: {traceback.format_exc()}")
                
                # Check if we should retry
                if self.state_manager.should_retry(session_id, ErrorType.SYSTEM_ERROR):
                    print("Retrying...")
                    self.state_manager.increment_retry(session_id)
                    return "Let me try that again... " + self.generate_response(question, session_id)
                else:
                    print("Max retries reached")
                    return "I'm having trouble answering your question. Please try asking something else or rephrase your question."
        
        except Exception as e:
            print(f"Complete failure in generate_response: {str(e)}")
            print(f"Complete traceback: {traceback.format_exc()}")
            return "I'm experiencing technical difficulties. Please try your question again."

class N8nConfig:
    def __init__(self):
        self.webhook_url = "http://localhost:5678/webhook-test/5879c0d0-def8-4d81-bbe7-adf19b68eebb"
        self.enabled = True  
        self.timeout = 5 
        self.retries = 2

n8n_config = N8nConfig()

def send_to_n8n_async(payload: dict):
    def _send():
        for attempt in range(n8n_config.retries + 1):
            try:
                response = requests.post(
                    n8n_config.webhook_url,
                    json=payload,
                    timeout=n8n_config.timeout
                )
                if response.status_code == 200:
                    print("Successfully sent to n8n")
                    break
                else:
                    print(f"n8n returned status {response.status_code}")
            except Exception as e:
                print(f"Attempt {attempt + 1} failed: {str(e)}")
                if attempt == n8n_config.retries:
                    print("Failed to send to n8n after retries")

    if n8n_config.enabled:
        Thread(target=_send).start()

# Initialize the system
try:
    print("Initializing RAG System...")
    rag_system = AgenticRAG()
    print("RAG System initialized successfully!")
    
    # Test function
    def ask_question(question: str, session_id: str = "default") -> str:
        """User-friendly function to ask questions"""
        return rag_system.generate_response(question, session_id)
    
    # Example usage
    # if __name__ == "__main__":
    #     # Test the system
    #     print("\n" + "="*50)
    #     print("TESTING RAG SYSTEM")
    #     print("="*50)
        
    #     test_questions = [
    #         "What is Assessli?",
    #         "Where is Assessli located?",
    #         "How can I contact Assessli?",
    #         "What services does Assessli provide?"
    #     ]
        
    #     for i, test_question in enumerate(test_questions, 1):
    #         print(f"\nTest {i}: {test_question}")
    #         print("-" * 50)
            
    #         try:
    #             response = ask_question(test_question, f"test_session")
    #             print(f"Response: {response}")
    #         except Exception as e:
    #             print(f"Test failed: {str(e)}")
    #             print(f"Traceback: {traceback.format_exc()}")
        
    #     print("\n" + "="*50)
    #     print("TESTING COMPLETE")
    #     print("="*50)
        
except Exception as e:
    print(f"Failed to initialize RAG system: {str(e)}")
    print(f"Initialization traceback: {traceback.format_exc()}")
    print("Please check your API keys and network connection.")
    
    # Create a minimal fallback system
    class FallbackRAGSystem:
        def generate_response(self, question: str, session_id: str = "default") -> str:
            return "The RAG system is currently unavailable. Please check your configuration and try again."
    
    rag_system = FallbackRAGSystem()
    
    def ask_question(question: str, session_id: str = "default") -> str:
        return rag_system.generate_response(question, session_id)
    

from flask import Flask, request, jsonify

app = Flask(__name__)

UPLOAD_FOLDER = tempfile.mkdtemp()
MAX_CONTENT_LENGTH = 50 * 1024 * 1024  # 50MB max file size
ALLOWED_EXTENSIONS = {
    'pdf', 'docx', 'doc', 'csv', 'json', 'xml', 'txt',
    'jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff',
    'mp4', 'avi', 'mov', 'wmv', 'flv', 'webm',
    'mp3', 'wav', 'aac', 'flac', 'ogg', 'wma'
}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_content_type(filename):
    """Determine content type based on file extension"""
    ext = filename.rsplit('.', 1)[1].lower()
    
    type_mapping = {
        'pdf': 'pdf',
        'docx': 'docx', 'doc': 'docx',
        'csv': 'csv',
        'json': 'json',
        'xml': 'xml',
        'txt': 'text',
        'jpg': 'image', 'jpeg': 'image', 'png': 'image', 
        'gif': 'image', 'bmp': 'image', 'tiff': 'image',
        'mp4': 'video', 'avi': 'video', 'mov': 'video', 
        'wmv': 'video', 'flv': 'video', 'webm': 'video',
        'mp3': 'audio', 'wav': 'audio', 'aac': 'audio', 
        'flac': 'audio', 'ogg': 'audio', 'wma': 'audio'
    }
    
    return type_mapping.get(ext, 'unknown')

def process_multimodal_content(file_path, content_type, text_query=None):
    try:
        content_data = rag_system.multimodal_processor.process_content(file_path, content_type)
        
        if 'error' in content_data:
            return f"Error processing file: {content_data['error']}"
        
        extracted_text = ""
        if content_type == 'pdf':
            extracted_text = content_data.get('text', '')
        elif content_type == 'docx':
            extracted_text = content_data.get('text', '')
        elif content_type == 'csv':
            summary = content_data.get('summary', {})
            extracted_text = f"CSV Data Summary: {len(content_data.get('data', []))} rows, {len(content_data.get('columns', []))} columns. Columns: {', '.join(content_data.get('columns', []))}"
        elif content_type == 'json':
            extracted_text = f"JSON Structure: {content_data.get('structure', {})}"
        elif content_type == 'image':
            extracted_text = content_data.get('text', 'No text found in image')
        elif content_type == 'video':
            extracted_text = f"Video metadata: Duration: {content_data.get('duration', 0):.2f}s, Resolution: {content_data.get('width', 0)}x{content_data.get('height', 0)}"
        elif content_type == 'audio':
            extracted_text = f"Audio metadata: Duration: {content_data.get('duration', 0):.2f}s, Sample Rate: {content_data.get('sample_rate', 0)}Hz"
        
        if text_query:
            combined_query = f"Based on the uploaded {content_type} content: {extracted_text[:1000]}... Please answer: {text_query}"
        else:
            combined_query = f"Please analyze and summarize this {content_type} content: {extracted_text[:1000]}..."
        
        return combined_query, content_data
        
    except Exception as e:
        return f"Error processing multimodal content: {str(e)}", None
    
def handle_multimodal_request():
    """Handle multimodal requests with files and text"""
    try:
        session_id = request.form.get('session_id', 'default')
        text_query = request.form.get('input', '')
        
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        if not allowed_file(file.filename):
            return jsonify({"error": f"File type not supported. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"}), 400
        
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        content_type = get_content_type(filename)
        
        result = process_multimodal_content(file_path, content_type, text_query)
        
        if isinstance(result, tuple):
            combined_query, content_data = result
        else:
            combined_query = result
            content_data = None
        
        # Analyze sentiment
        sentiment = rag_system.analyze_sentiment(combined_query, session_id)
        trend_label, trend_score = rag_system.get_sentiment_trend(session_id)
        
        # Generate response
        output = ask_question(combined_query, session_id)
        
        response_data = {
            "output": output,
            "file_info": {
                "filename": filename,
                "content_type": content_type,
                "processed": content_data is not None
            }
        }
        
        if content_data and not content_data.get('error'):
            if content_type == 'csv':
                response_data["file_info"]["rows"] = len(content_data.get('data', []))
                response_data["file_info"]["columns"] = content_data.get('columns', [])
            elif content_type == 'image':
                response_data["file_info"]["has_text"] = content_data.get('has_text', False)
            elif content_type == 'video':
                response_data["file_info"]["duration"] = content_data.get('duration', 0)
            elif content_type == 'audio':
                response_data["file_info"]["duration"] = content_data.get('duration', 0)
        
        # Send to n8n
        n8n_payload = {
            'question': text_query,
            'response': output,
            'sentiment': {
                'label': sentiment.label.value,
                'score': sentiment.score
            },
            'session_trend': {
                'label': trend_label.value,
                'score': trend_score
            },
            'timestamp': datetime.now().isoformat(),
            'session_id': session_id,
            'content_type': 'multimodal'
        }
        
        send_to_n8n_async(n8n_payload)
        
        return jsonify(response_data)
            

                
    except Exception as e:
        return jsonify({"error": f"Multimodal processing failed: {str(e)}"}), 500


@app.route('/chat', methods=['POST'])
def predict():
    try:
        if request.content_type and request.content_type.startswith('multipart/form-data'):
            return handle_multimodal_request()
        
        data = request.json
        if 'input' not in data:
            return jsonify({"error": "No input data provided"}), 400
        
        session_id = data.get('session_id', 'default')
        question = data['input']
        
        sentiment = rag_system.analyze_sentiment(question, session_id)
        trend_label, trend_score = rag_system.get_sentiment_trend(session_id)
        
        output = ask_question(data['input'], data['session_id'])
        
        n8n_payload = {
            'question': question,
            'response': output,
            'sentiment': {
                'label': sentiment.label.value,
                'score': sentiment.score
            },
            'session_trend': {
                'label': trend_label.value,
                'score': trend_score
            },
            'timestamp': datetime.now().isoformat(),
            'session_id': session_id,
        }
        
        send_to_n8n_async(n8n_payload)
        
        return jsonify({"output": output})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.errorhandler(requests.exceptions.RequestException)
def handle_n8n_error(e):
    print(f"n8n communication error: {str(e)}")
    return None

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
    
